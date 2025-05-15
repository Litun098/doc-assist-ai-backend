"""
Document service for managing documents with Supabase.
"""
import os
import uuid
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional
from fastapi import UploadFile, HTTPException, BackgroundTasks
from supabase import create_client, Client

from app.utils.s3_storage import s3_storage
from app.services.document_processor import document_processor
from config.config import settings

# Configure logging
logger = logging.getLogger(__name__)

# Import connection manager
from app.utils.connection_manager import connection_manager

# Initialize Supabase client using connection manager
try:
    # Get Supabase client from connection manager
    supabase = connection_manager.get_supabase_client("default")
    if supabase:
        logger.info(f"Connected to Supabase at {settings.SUPABASE_URL}")
    else:
        logger.error("Failed to get Supabase client from connection manager")
except Exception as e:
    logger.error(f"Error connecting to Supabase: {str(e)}")
    supabase = None

class DocumentService:
    """Document service for managing documents with Supabase."""

    def __init__(self):
        """Initialize the document service."""
        self.supabase = supabase

    async def upload_document(self, file: UploadFile, user_id: str, background_tasks: BackgroundTasks) -> Dict[str, Any]:
        """
        Upload a document and start processing.

        Args:
            file: The file to upload
            user_id: ID of the user
            background_tasks: FastAPI background tasks

        Returns:
            Document information
        """
        try:
            # Check file extension
            file_ext = file.filename.split('.')[-1].lower()
            if file_ext not in settings.ALLOWED_EXTENSIONS:
                raise HTTPException(
                    status_code=400,
                    detail=f"File type not allowed. Allowed types: {', '.join(settings.ALLOWED_EXTENSIONS)}"
                )

            # Check file size
            file.file.seek(0, 2)  # Seek to end
            file_size = file.file.tell()
            file.file.seek(0)  # Reset file position

            if file_size > settings.MAX_UPLOAD_SIZE:
                raise HTTPException(
                    status_code=400,
                    detail=f"File too large. Maximum size: {settings.MAX_UPLOAD_SIZE / (1024 * 1024):.1f} MB"
                )

            # Create a unique file ID
            file_id = str(uuid.uuid4())

            # Upload to S3 if available, otherwise local storage
            if s3_storage.is_available():
                s3_key = f"users/{user_id}/documents/{file_id}.{file_ext}"
                s3_result = await s3_storage.upload_file(file, s3_key)
                file_url = s3_result["url"]
                storage_type = "s3"
            else:
                # Fallback to local storage
                os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
                file_path = os.path.join(settings.UPLOAD_DIR, f"{file_id}.{file_ext}")
                file_content = await file.read()
                with open(file_path, "wb") as f:
                    f.write(file_content)
                file_url = file_path
                storage_type = "local"
                s3_key = file_path  # Use local path as key

            # Create document record in Supabase if available
            document_data = {
                "id": file_id,
                "user_id": user_id,
                "file_name": file.filename,
                "file_type": file_ext,
                "file_size": file_size,
                "s3_key": s3_key,
                "status": "processing",
                "created_at": datetime.now().isoformat(),
                "updated_at": datetime.now().isoformat()
            }

            if self.supabase:
                try:
                    # Try using service role key first to avoid RLS issues
                    if settings.SUPABASE_SERVICE_KEY:
                        try:
                            logger.info(f"Inserting document data using service role for user ID: {user_id}")
                            service_supabase = create_client(
                                supabase_url=settings.SUPABASE_URL,
                                supabase_key=settings.SUPABASE_SERVICE_KEY
                            )
                            service_supabase.table("documents").insert(document_data).execute()
                            logger.info(f"Document data inserted successfully using service role for file ID: {file_id}")
                        except Exception as service_error:
                            logger.error(f"Error inserting document data using service role: {str(service_error)}")
                            # Fall back to regular key
                            logger.info(f"Falling back to regular key for inserting document data for user ID: {user_id}")
                            self.supabase.table("documents").insert(document_data).execute()
                            logger.info(f"Document data inserted successfully for file ID: {file_id}")
                    else:
                        # No service key available, use regular key
                        logger.info(f"Inserting document data into Supabase for user ID: {user_id}")
                        self.supabase.table("documents").insert(document_data).execute()
                        logger.info(f"Document data inserted successfully for file ID: {file_id}")
                except Exception as insert_error:
                    logger.error(f"Error inserting document data: {str(insert_error)}")
                    # Continue with document processing despite the error
                    logger.info(f"Continuing with document processing despite database insert error for file ID: {file_id}")

            # Start document processing in background
            background_tasks.add_task(
                self._process_document,
                file_url=file_url,
                file_id=file_id,
                user_id=user_id,
                storage_type=storage_type
            )

            return {
                "file_id": file_id,
                "file_name": file.filename,
                "file_type": file_ext,
                "status": "processing",
                "message": f"Document uploaded to {storage_type} storage and processing started"
            }

        except HTTPException as e:
            raise e
        except Exception as e:
            import traceback
            logger.error(f"Error uploading document: {str(e)}")
            logger.error(f"Traceback: {traceback.format_exc()}")

            # Check if it's an RLS policy error
            error_str = str(e)
            if "row-level security policy" in error_str and "documents" in error_str:
                logger.warning("RLS policy error detected for documents table")
                logger.warning("Please set up proper RLS policies for the documents table")

                # Instead of using mock data, we'll use the service role to bypass RLS
                # This ensures we're still using the actual database
                try:
                    # Use service role to insert document data (bypasses RLS)
                    service_supabase = create_client(
                        supabase_url=settings.SUPABASE_URL,
                        supabase_key=settings.SUPABASE_SERVICE_KEY
                    )

                    logger.info(f"Attempting to insert document data using service role for user ID: {user_id}")
                    service_supabase.table("documents").insert(document_data).execute()
                    logger.info(f"Document data inserted successfully using service role for file ID: {file_id}")

                    # Continue with normal flow
                    return {
                        "file_id": file_id,
                        "file_name": file.filename,
                        "file_type": file_ext,
                        "status": "processing",
                        "message": f"Document uploaded to {storage_type} storage and processing started"
                    }
                except Exception as service_error:
                    logger.error(f"Error inserting document data using service role: {str(service_error)}")
                    # If service role also fails, we'll fall through to the general error handler
                    raise HTTPException(
                        status_code=500,
                        detail=f"Error uploading document: {str(service_error)}"
                    )

            raise HTTPException(
                status_code=500,
                detail=f"Error uploading document: {str(e)}"
            )

    async def list_documents(self, user_id: str) -> Dict[str, List[Dict[str, Any]]]:
        """
        List all documents for a user.

        Args:
            user_id: ID of the user

        Returns:
            List of documents
        """
        try:
            documents = []

            # Get documents from Supabase if available
            if self.supabase:
                try:
                    # Try using service role key first to avoid RLS issues
                    if settings.SUPABASE_SERVICE_KEY:
                        try:
                            logger.info(f"Listing documents using service role for user ID: {user_id}")
                            service_supabase = create_client(
                                supabase_url=settings.SUPABASE_URL,
                                supabase_key=settings.SUPABASE_SERVICE_KEY
                            )
                            response = service_supabase.table("documents").select("*").eq("user_id", user_id).execute()
                            logger.info(f"Documents listed successfully using service role for user ID: {user_id}")
                        except Exception as service_error:
                            logger.error(f"Error listing documents using service role: {str(service_error)}")
                            # Fall back to regular key
                            logger.info(f"Falling back to regular key for listing documents for user ID: {user_id}")
                            response = self.supabase.table("documents").select("*").eq("user_id", user_id).execute()
                            logger.info(f"Documents listed successfully for user ID: {user_id}")
                    else:
                        # No service key available, use regular key
                        response = self.supabase.table("documents").select("*").eq("user_id", user_id).execute()

                    for doc in response.data:
                        documents.append({
                            "file_id": doc["id"],
                            "file_name": doc["file_name"],
                            "file_type": doc["file_type"],
                            "file_size": doc.get("file_size", 0),  # Include file size
                            "status": doc["status"],
                            "created_at": doc["created_at"]
                        })
                except Exception as list_error:
                    logger.error(f"Error listing documents: {str(list_error)}")
                    # Continue with empty documents list
            else:
                # Fallback to local storage
                if os.path.exists(settings.UPLOAD_DIR):
                    for filename in os.listdir(settings.UPLOAD_DIR):
                        file_path = os.path.join(settings.UPLOAD_DIR, filename)
                        if os.path.isfile(file_path):
                            file_id, ext = os.path.splitext(filename)
                            ext = ext.lstrip('.')

                            # Get file size
                            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

                            documents.append({
                                "file_id": file_id,
                                "file_name": filename,
                                "file_type": ext,
                                "file_size": file_size,
                                "status": "processed"
                            })

            return {"documents": documents}

        except Exception as e:
            logger.error(f"Error listing documents: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error listing documents: {str(e)}"
            )

    async def get_document_preview(self, file_id: str, user_id: str) -> Dict[str, Any]:
        """
        Get a preview of a document.

        Args:
            file_id: ID of the file to preview
            user_id: ID of the user

        Returns:
            Document preview data
        """
        try:
            document = None
            preview_text = ""

            # Check if document exists in Supabase
            if self.supabase:
                try:
                    # Try using service role key first to avoid RLS issues
                    if settings.SUPABASE_SERVICE_KEY:
                        try:
                            logger.info(f"Getting document using service role for file ID: {file_id}")
                            service_supabase = create_client(
                                supabase_url=settings.SUPABASE_URL,
                                supabase_key=settings.SUPABASE_SERVICE_KEY
                            )
                            response = service_supabase.table("documents").select("*").eq("id", file_id).eq("user_id", user_id).execute()
                            logger.info(f"Document retrieved successfully using service role for file ID: {file_id}")
                        except Exception as service_error:
                            logger.error(f"Error getting document using service role: {str(service_error)}")
                            # Fall back to regular key
                            logger.info(f"Falling back to regular key for getting document for file ID: {file_id}")
                            response = self.supabase.table("documents").select("*").eq("id", file_id).eq("user_id", user_id).execute()
                            logger.info(f"Document retrieved successfully for file ID: {file_id}")
                    else:
                        # No service key available, use regular key
                        response = self.supabase.table("documents").select("*").eq("id", file_id).eq("user_id", user_id).execute()

                    if not response.data:
                        raise HTTPException(
                            status_code=404,
                            detail=f"Document with ID {file_id} not found or does not belong to user"
                        )

                    document = response.data[0]

                    # Get document from S3 if available
                    if s3_storage.is_available() and document["s3_key"]:
                        # Get the file from S3
                        temp_file_path = os.path.join(settings.UPLOAD_DIR, f"temp_{file_id}")
                        s3_storage.download_file(document["s3_key"], temp_file_path)

                        # Extract preview text based on file type
                        file_type = document["file_type"].lower()
                        if file_type == "pdf":
                            # Extract first page of PDF
                            from pypdf import PdfReader
                            reader = PdfReader(temp_file_path)
                            if len(reader.pages) > 0:
                                preview_text = reader.pages[0].extract_text()[:1000]  # First 1000 chars
                        elif file_type in ["docx", "doc"]:
                            # Extract first page of Word document
                            import docx
                            doc = docx.Document(temp_file_path)
                            if len(doc.paragraphs) > 0:
                                preview_text = "\n".join([p.text for p in doc.paragraphs[:5]])[:1000]  # First 5 paragraphs
                        elif file_type in ["txt", "md"]:
                            # Extract first 1000 chars of text file
                            with open(temp_file_path, "r", encoding="utf-8") as f:
                                preview_text = f.read(1000)

                        # Clean up temp file
                        if os.path.exists(temp_file_path):
                            os.remove(temp_file_path)
                except Exception as e:
                    logger.error(f"Error getting document preview: {str(e)}")
                    preview_text = "Error generating preview. Please select the document to chat with it."
            else:
                # Fallback to local storage
                file_path = None
                file_name = None
                file_type = None

                if os.path.exists(settings.UPLOAD_DIR):
                    for filename in os.listdir(settings.UPLOAD_DIR):
                        if filename.startswith(file_id):
                            file_path = os.path.join(settings.UPLOAD_DIR, filename)
                            file_name = filename
                            _, ext = os.path.splitext(filename)
                            file_type = ext.lstrip('.')
                            break

                if not file_path:
                    raise HTTPException(
                        status_code=404,
                        detail=f"File with ID {file_id} not found"
                    )

                # Extract preview text based on file type
                if file_type.lower() == "pdf":
                    # Extract first page of PDF
                    from pypdf import PdfReader
                    reader = PdfReader(file_path)
                    if len(reader.pages) > 0:
                        preview_text = reader.pages[0].extract_text()[:1000]  # First 1000 chars
                elif file_type.lower() in ["docx", "doc"]:
                    # Extract first page of Word document
                    import docx
                    doc = docx.Document(file_path)
                    if len(doc.paragraphs) > 0:
                        preview_text = "\n".join([p.text for p in doc.paragraphs[:5]])[:1000]  # First 5 paragraphs
                elif file_type.lower() in ["txt", "md"]:
                    # Extract first 1000 chars of text file
                    with open(file_path, "r", encoding="utf-8") as f:
                        preview_text = f.read(1000)

                document = {
                    "id": file_id,
                    "file_name": file_name,
                    "file_type": file_type
                }

            return {
                "file_id": file_id,
                "file_name": document["file_name"],
                "file_type": document["file_type"],
                "preview_text": preview_text or "No preview available. Please select the document to chat with it."
            }

        except HTTPException as e:
            raise e
        except Exception as e:
            logger.error(f"Error getting document preview: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error getting document preview: {str(e)}"
            )

    async def delete_document(self, file_id: str, user_id: str) -> Dict[str, Any]:
        """
        Delete a document.

        Args:
            file_id: ID of the file to delete
            user_id: ID of the user

        Returns:
            Deletion status
        """
        try:
            document = None

            # Check if document exists in Supabase
            if self.supabase:
                response = self.supabase.table("documents").select("*").eq("id", file_id).eq("user_id", user_id).execute()

                if not response.data:
                    raise HTTPException(
                        status_code=404,
                        detail=f"Document with ID {file_id} not found or does not belong to user"
                    )

                document = response.data[0]

                # Delete from S3 if available
                if s3_storage.is_available() and document["s3_key"].startswith("users/"):
                    s3_storage.delete_file(document["s3_key"])

                # Delete from Supabase
                self.supabase.table("documents").delete().eq("id", file_id).execute()
            else:
                # Fallback to local storage
                file_path = None
                file_name = None
                file_type = None

                if os.path.exists(settings.UPLOAD_DIR):
                    for filename in os.listdir(settings.UPLOAD_DIR):
                        if filename.startswith(file_id):
                            file_path = os.path.join(settings.UPLOAD_DIR, filename)
                            file_name = filename
                            _, ext = os.path.splitext(filename)
                            file_type = ext.lstrip('.')
                            break

                if not file_path:
                    raise HTTPException(
                        status_code=404,
                        detail=f"File with ID {file_id} not found"
                    )

                # Delete the file
                os.remove(file_path)

                document = {
                    "file_name": file_name,
                    "file_type": file_type
                }

            return {
                "file_id": file_id,
                "file_name": document["file_name"],
                "file_type": document["file_type"],
                "status": "deleted",
                "message": "Document deleted successfully"
            }

        except HTTPException as e:
            raise e
        except Exception as e:
            logger.error(f"Error deleting document: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error deleting document: {str(e)}"
            )

    async def _process_document(self, file_url: str, file_id: str, user_id: str, storage_type: str):
        """
        Process a document in the background.

        Args:
            file_url: URL or path to the file
            file_id: ID of the file
            user_id: ID of the user
            storage_type: Type of storage ('s3' or 'local')
        """
        # For large documents, use Celery task instead of processing directly
        # This helps prevent timeouts and memory issues
        file_size = 0
        file_ext = ""

        # Get file size and type
        if storage_type == "local" and os.path.exists(file_url):
            file_size = os.path.getsize(file_url)
            file_ext = os.path.splitext(file_url)[1].lstrip('.')
        elif storage_type == "s3":
            # For S3, we need to get the file info from Supabase
            if self.supabase:
                try:
                    # Try using service role key first to avoid RLS issues
                    if settings.SUPABASE_SERVICE_KEY:
                        try:
                            logger.info(f"Getting document info using service role for file ID: {file_id}")
                            service_supabase = create_client(
                                supabase_url=settings.SUPABASE_URL,
                                supabase_key=settings.SUPABASE_SERVICE_KEY
                            )
                            response = service_supabase.table("documents").select("*").eq("id", file_id).execute()
                            if response.data:
                                file_size = response.data[0].get("file_size", 0)
                                file_ext = response.data[0].get("file_type", "")
                        except Exception as e:
                            logger.error(f"Error getting document info using service role: {str(e)}")
                except Exception as e:
                    logger.error(f"Error getting document info: {str(e)}")

        # Use Celery for large files (> 5MB) or if file size is unknown
        large_file_threshold = 5 * 1024 * 1024  # 5MB
        if file_size > large_file_threshold or file_size == 0:
            logger.info(f"Using Celery task for large document processing: {file_id} ({file_size} bytes)")
            # Import here to avoid circular imports
            from app.workers.llama_index_tasks import process_file_with_llama_index

            # Start Celery task
            process_file_with_llama_index.delay(
                file_id=file_id,
                user_id=user_id,
                file_path=file_url,
                file_type=file_ext
            )
            return
        try:
            # Process the document
            result = document_processor.process_document(
                file_path=file_url,
                file_id=file_id,
                user_id=user_id,
                storage_type=storage_type
            )

            # Update document status in Supabase
            if self.supabase and result["status"] == "success":
                try:
                    # Try using service role key first to avoid RLS issues
                    if settings.SUPABASE_SERVICE_KEY:
                        try:
                            logger.info(f"Updating document status to processed using service role for file ID: {file_id}")
                            service_supabase = create_client(
                                supabase_url=settings.SUPABASE_URL,
                                supabase_key=settings.SUPABASE_SERVICE_KEY
                            )
                            service_supabase.table("documents").update({
                                "status": "processed",
                                "processed_at": datetime.now().isoformat(),
                                "updated_at": datetime.now().isoformat(),
                                "metadata": {
                                    "num_chunks": result.get("num_chunks", 0),
                                    "chunking_strategy": result.get("chunking_strategy", "unknown")
                                }
                            }).eq("id", file_id).execute()
                            logger.info(f"Document status updated successfully using service role for file ID: {file_id}")
                        except Exception as service_error:
                            logger.error(f"Error updating document status using service role: {str(service_error)}")
                            # Fall back to regular key
                            logger.info(f"Falling back to regular key for updating document status for file ID: {file_id}")
                            self.supabase.table("documents").update({
                                "status": "processed",
                                "processed_at": datetime.now().isoformat(),
                                "updated_at": datetime.now().isoformat(),
                                "metadata": {
                                    "num_chunks": result.get("num_chunks", 0),
                                    "chunking_strategy": result.get("chunking_strategy", "unknown")
                                }
                            }).eq("id", file_id).execute()
                            logger.info(f"Document status updated successfully for file ID: {file_id}")
                    else:
                        # No service key available, use regular key
                        logger.info(f"Updating document status to processed for file ID: {file_id}")
                        self.supabase.table("documents").update({
                            "status": "processed",
                            "processed_at": datetime.now().isoformat(),
                            "updated_at": datetime.now().isoformat(),
                            "metadata": {
                                "num_chunks": result.get("num_chunks", 0),
                                "chunking_strategy": result.get("chunking_strategy", "unknown")
                            }
                        }).eq("id", file_id).execute()
                        logger.info(f"Document status updated successfully for file ID: {file_id}")
                except Exception as update_error:
                    logger.error(f"Error updating document status: {str(update_error)}")
                    # Continue despite the error
                    logger.info(f"Document was processed successfully despite database update error for file ID: {file_id}")
            elif self.supabase and result["status"] == "error":
                try:
                    # Try using service role key first to avoid RLS issues
                    if settings.SUPABASE_SERVICE_KEY:
                        try:
                            logger.info(f"Updating document status to error using service role for file ID: {file_id}")
                            service_supabase = create_client(
                                supabase_url=settings.SUPABASE_URL,
                                supabase_key=settings.SUPABASE_SERVICE_KEY
                            )
                            service_supabase.table("documents").update({
                                "status": "error",
                                "error_message": result.get("error", "Unknown error"),
                                "updated_at": datetime.now().isoformat()
                            }).eq("id", file_id).execute()
                            logger.info(f"Document error status updated successfully using service role for file ID: {file_id}")
                        except Exception as service_error:
                            logger.error(f"Error updating document error status using service role: {str(service_error)}")
                            # Fall back to regular key
                            logger.info(f"Falling back to regular key for updating document error status for file ID: {file_id}")
                            self.supabase.table("documents").update({
                                "status": "error",
                                "error_message": result.get("error", "Unknown error"),
                                "updated_at": datetime.now().isoformat()
                            }).eq("id", file_id).execute()
                            logger.info(f"Document error status updated successfully for file ID: {file_id}")
                    else:
                        # No service key available, use regular key
                        logger.info(f"Updating document status to error for file ID: {file_id}")
                        self.supabase.table("documents").update({
                            "status": "error",
                            "error_message": result.get("error", "Unknown error"),
                            "updated_at": datetime.now().isoformat()
                        }).eq("id", file_id).execute()
                        logger.info(f"Document error status updated successfully for file ID: {file_id}")
                except Exception as update_error:
                    logger.error(f"Error updating document error status: {str(update_error)}")
                    # Continue despite the error

        except Exception as e:
            logger.error(f"Error processing document {file_id}: {str(e)}")
            # Update document status in Supabase
            if self.supabase:
                try:
                    # Try using service role key first to avoid RLS issues
                    if settings.SUPABASE_SERVICE_KEY:
                        try:
                            logger.info(f"Updating document status to error using service role after exception for file ID: {file_id}")
                            service_supabase = create_client(
                                supabase_url=settings.SUPABASE_URL,
                                supabase_key=settings.SUPABASE_SERVICE_KEY
                            )
                            service_supabase.table("documents").update({
                                "status": "error",
                                "error_message": str(e),
                                "updated_at": datetime.now().isoformat()
                            }).eq("id", file_id).execute()
                            logger.info(f"Document error status updated successfully using service role after exception for file ID: {file_id}")
                        except Exception as service_error:
                            logger.error(f"Error updating document error status using service role after exception: {str(service_error)}")
                            # Fall back to regular key
                            logger.info(f"Falling back to regular key for updating document error status after exception for file ID: {file_id}")
                            self.supabase.table("documents").update({
                                "status": "error",
                                "error_message": str(e),
                                "updated_at": datetime.now().isoformat()
                            }).eq("id", file_id).execute()
                            logger.info(f"Document error status updated successfully after exception for file ID: {file_id}")
                    else:
                        # No service key available, use regular key
                        logger.info(f"Updating document status to error after exception for file ID: {file_id}")
                        self.supabase.table("documents").update({
                            "status": "error",
                            "error_message": str(e),
                            "updated_at": datetime.now().isoformat()
                        }).eq("id", file_id).execute()
                        logger.info(f"Document error status updated successfully after exception for file ID: {file_id}")
                except Exception as update_error:
                    logger.error(f"Error updating document error status after exception: {str(update_error)}")
                    # Continue despite the error

# Create a singleton instance
document_service = DocumentService()
